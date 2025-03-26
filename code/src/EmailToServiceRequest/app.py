from flask import Flask, request, jsonify
import os
from werkzeug.utils import secure_filename
from docx import Document
import pandas as pd
import PyPDF2
import fitz  # PyMuPDF
from PIL import Image
import pytesseract
import magic
import io
from email import message_from_file
from email.policy import default
from email.utils import parseaddr
from llama_cpp import Llama
from gpt4all import GPT4All 
from collections import Counter
import json

app = Flask(__name__)

pytesseract.pytesseract.tesseract_cmd = "C:\\Program Files\\Tesseract-OCR\\tesseract.exe"

# Define allowed file types
ALLOWED_ATTACHMENT_EXTENSIONS = {'pdf', 'doc', 'docx', 'xls', 'xlsx', 'txt', 'csv', 'jpg', 'jpeg', 'png'}
ALLOWED_EMIAL_EXTENSIONS = {'eml'}

# Ensure you have a folder for uploads
UPLOAD_FOLDER = 'uploads'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

def allowed_attachment_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_ATTACHMENT_EXTENSIONS

def allowed_email_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EMIAL_EXTENSIONS

# Helper function to read different file types
def read_pdf(file_path):
    # Open the PDF
    doc = fitz.open(file_path)
    text = ''
    
    # Extract text from PDF pages
    for page_num in range(doc.page_count):
        page = doc.load_page(page_num)
        text += page.get_text()
        
        # Extract images and perform OCR on them
        for img_index in range(len(page.get_images(full=True))):
            xref = page.get_images(full=True)[img_index][0]
            pix = fitz.Pixmap(doc, xref)
            if pix.n < 4:  # if the image is not CMYK
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                ocr_text = pytesseract.image_to_string(img)
                text += ocr_text

    return text

def read_docx(file_path):
    doc = Document(file_path)
    text = ''
    
    # Extract text from the paragraphs
    for para in doc.paragraphs:
        text += para.text + '\n'
        
    # Handle images inside the DOCX file and apply OCR on them
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image = Image.open(io.BytesIO(rel.target_part.blob))
            text += pytesseract.image_to_string(image)
    
    return text

def read_txt(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()

def read_csv(file_path):
    df = pd.read_csv(file_path)
    return df.to_string()

def read_image(file_path):
    img = Image.open(file_path)
    return pytesseract.image_to_string(img)

def read_xls(file_path):
    df = pd.read_excel(file_path)
    return df.to_string()

def processEmailBody(part):
    body = ""
    body_plain_or_html = ""
    body_octo = ""
    body_multi = ""

    content_type = part.get_content_type()
    
    if content_type in ["text/plain", "text/html"]:
        body_plain_or_html = part.get_payload(decode=True).decode(part.get_content_charset(), 'ignore')
    elif content_type.startswith("image/"):
        # Extract image body and apply OCR
        img_data = part.get_payload(decode=True)
        img = Image.open(io.BytesIO(img_data))
        ocr_text = pytesseract.image_to_string(img)
        body_octo += ocr_text.encode('utf-8')
    elif content_type == "multipart/alternative":
        bParts = part.iter_parts()
        for bp in bParts:
            bpBody = processEmailBody(bp)
            body_multi += bpBody +'\n'
                
    if body_plain_or_html:
        body += body_plain_or_html + '\n'
    if body_octo:
        body += body_octo + '\n'
    if body_multi:
        body += body_multi+ '\n'

    return body
        
def parse_eml(file_path):
    with open(file_path, 'r') as f:
        msg = message_from_file(f, policy=default)

    # Extract basic email fields
    email_info = {
        'From': msg['From'],
        'To': msg['To'],
        'CC': msg['CC'],
        'BCC': msg['BCC'],
        'Subject': msg['Subject'],
        'Body': "",
        'Attachments': []
    }

    # Extract the body of the email
    if msg.is_multipart():
        for part in msg.iter_parts():
            content_disposition = str(part.get("Content-Disposition"))
            if "attachment" not in content_disposition:
                email_info['Body'] = processEmailBody(part)
            else:
                filename = part.get_filename()
                attach_file_path = os.path.join(app.config['UPLOAD_FOLDER'], 'AttachTemp', filename)
                if os.path.exists(attach_file_path):
                    os.remove(attach_file_path)
                with open(attach_file_path, "wb") as f:
                    f.write(io.BytesIO(part.get_payload(decode=True)).getbuffer())
                    f.close()
                file_type = magic.from_file(attach_file_path, mime=True)
                if filename and allowed_attachment_file(filename):
                    attachment_info = {
                        'filename': filename,
                        'file_content': part.get_payload(decode=True),
                        'file_type': file_type
                    }
                    email_info['Attachments'].append(attachment_info)
        
    else:
        email_info['Body'] = msg.get_payload(decode=True).decode(msg.get_content_charset(), 'ignore')

    return email_info

def parse_eml_and_attachments(file_path):
    email_info = parse_eml(file_path)

    # Process each attachment if present
    for attachment in email_info['Attachments']:
        file_data = io.BytesIO(attachment['file_content'])
        attach_file_path = os.path.join(app.config['UPLOAD_FOLDER'], 'AttachTemp', attachment['filename'])
        # if os.path.exists(attach_file_path):
        #     os.remove(attach_file_path)
        # with open(attach_file_path, "wb") as f:
        #     f.write(file_data.getbuffer())
        
        attachment['extracted_text'] = ""
        if attachment['file_type'] == 'application/pdf':
            attachment['extracted_text'] = read_pdf(attach_file_path)
        elif attachment['file_type'] in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']:
            attachment['extracted_text'] = read_docx(attach_file_path)
        elif attachment['file_type'] == 'text/plain':
            attachment['extracted_text'] = read_txt(attach_file_path)
        elif attachment['file_type'] == 'application/vnd.ms-excel' or attachment['file_type'] == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':
            attachment['extracted_text'] = read_xls(attach_file_path)
        elif attachment['file_type'] == 'text/csv':
            attachment['extracted_text'] = read_csv(attach_file_path)
        elif attachment['file_type'].startswith("image/"):
            attachment['extracted_text'] = read_image(attach_file_path)
        
        attachment['file_content'] = None

        if os.path.exists(attach_file_path):
            os.remove(attach_file_path)

    return email_info
def read_docx_no_paragraphs(file_path):
    doc=Document(file_path)
    return " ".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])

def split_text(text,max_length=400):
    return [text[i:i+max_length]for i in range(0,len(text),max_length)]

@app.route('/process/file', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({"error": "No file part"}), 400

    file = request.files['file']

    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400

    if file and allowed_email_file(file.filename):
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)

        try:
            email_info = parse_eml_and_attachments(file_path)

            if os.path.exists(file_path):
                os.remove(file_path)
            model_file="Meta-Llama-3-8B-Instruct-Q4_0.gguf"
            model_folder = r"C:\Users\padhideb"
            model_path = os.path.join(model_folder,model_file)
            # file_name="Invoice.docx"
            # file_path = os.path.join(model_folder,file_name)
            llm=Llama(model_path)
            document_text = read_docx_no_paragraphs(email_info.body)
            chunks = split_text(document_text,max_length=400)

            few_shot_prompt = (
                "You are an expert at classifying documents by their request type."
                "Based on the content, assign one of the following labels:Finacial, Techincal, or Legal.\n\n"
                "Return only Answer with no explanation.\n\n"
                "Example 1:\n"
                "Document:Reason for transferring funds\n"
                "Request Type:Financial\n\n"
                "Example 2:\n"
                "Document:API integration guide explaining authentication procedures.\n"
                "Request Type:Technical\n\n"
                "Now classify the following document.\n"
                "Document:{document_chunk}\n"
                "Request Type:"
            )
            results = []
            for chunk in chunks:
                prompt = few_shot_prompt.format(document_chunk=chunk)
                response=llm(prompt)
                answer = response["choices"][0]["text"].strip()
                results.append(answer)

            final_classification = Counter(results).most_common(1)[0][0]
            output = {"request_type":final_classification}

            return jsonify({"email_info": email_info}), 200
        except Exception as e:
            return jsonify({"error": "Invalid file format"}), 400
    

@app.route('/process/path', methods=['POST'])
def upload_eml_directory():
    # Get the directory path from the request
    directory_path = request.json.get("directory_path")
    
    if not directory_path or not os.path.exists(directory_path) or not os.path.isdir(directory_path):
        return jsonify({"error": "Invalid directory path"}), 400

    email_info_list = []
    
    try:
        # List all .eml files in the given directory
        eml_files = [f for f in os.listdir(directory_path) if f.endswith('.eml')]

        for eml_file in eml_files:
            eml_file_path = os.path.join(directory_path, eml_file)
            email_info = parse_eml_and_attachments(eml_file_path)
            email_info_list.append(email_info)
        
        # Return a JSON array of email info objects
        return jsonify({"emails": email_info_list}), 200

    except Exception as e:
        return jsonify({"error": f"An error occurred while processing the directory: {str(e)}"}), 500


if __name__ == '__main__':
    app.run(host='127.0.0.1', port=5000,debug=True)
